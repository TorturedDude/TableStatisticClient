from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER

def CreateWordDocument(text,path,name_of_file):
    document = Document()
    p_header = document.add_paragraph()
    p_header.add_run('Договор').bold = True
    p_header.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = document.add_paragraph(text)

    global fullpath

    if path == '':
        fullpath = name_of_file + '.docx'
    else:
        fullpath = str(path) + '/' + str(name_of_file) + '.docx'

    document.save(fullpath)